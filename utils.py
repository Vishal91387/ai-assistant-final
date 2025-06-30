from io import BytesIO
from fpdf import FPDF
from docx import Document
import re

def clean_text(text):
    return re.sub(r'[^\x00-\x7F]+', '', text)

def export_chat_to_pdf(chat_log, summary=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    # Chat log
    for i, (question, answer) in enumerate(chat_log, start=1):
        q = clean_text(f"Q{i}: {question}")
        a = clean_text(f"A{i}: {answer}")
        pdf.multi_cell(0, 10, q, align="L")
        pdf.multi_cell(0, 10, a, align="L")
        pdf.ln()

    # Add summary if provided
    if summary:
        pdf.set_font("Arial", style="B", size=12)
        pdf.cell(0, 10, "Summary:", ln=True)
        pdf.set_font("Arial", size=12)
        for line in summary.strip().split("\n"):
            pdf.multi_cell(0, 10, clean_text(line), align="L")
        pdf.ln()

    buffer = BytesIO()
    pdf_bytes = pdf.output(dest='S').encode('latin-1', 'replace')
    buffer.write(pdf_bytes)
    buffer.seek(0)
    return buffer

def export_chat_to_word(chat_log, summary=None):
    doc = Document()
    doc.add_heading("Chat Log", 0)

    for i, (question, answer) in enumerate(chat_log, start=1):
        doc.add_paragraph(f"Q{i}: {question}", style="List Number")
        doc.add_paragraph(f"A{i}: {answer}")

    if summary:
        doc.add_heading("Summary", level=1)
        for line in summary.strip().split("\n"):
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
def export_chat_to_text(chat_log):
    buffer = BytesIO()
    content = ""
    for i, (question, answer) in enumerate(chat_log, start=1):
        content += f"Q{i}: {question}\nA{i}: {answer}\n\n"
    buffer.write(content.encode("utf-8"))
    buffer.seek(0)
    return buffer

# -----------------------------
# 🔍 Summarize Answer Helper
# -----------------------------
def summarize_answer(answer):
    import re
    from langchain_core.output_parsers import StrOutputParser
    from langchain_core.prompts import PromptTemplate
    from langchain_community.llms import Ollama

    llm = Ollama(model="llama3:8b")

    prompt = PromptTemplate.from_template(
    """You are an assistant that summarizes factual answers into short, crisp bullet points.

Please convert the following answer into 3 to 5 concise bullet points. Each point must start with a dash (-), and only include essential facts. Avoid repetition.

Answer:
{answer}

Bullet Point Summary:
- """
)
    chain = prompt | llm | StrOutputParser()

    try:
        result = chain.invoke({"answer": answer})
        print("🔍 SUMMARY RESULT:", result)  # Debug print
        points = re.findall(r"-\s+(.*)", result)
        return points if points else ["No summary available."]
    except Exception as e:
        print("❌ Summary generation failed:", e)
        return ["Summary failed."]